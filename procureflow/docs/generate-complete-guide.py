#!/usr/bin/env python3
"""
Generate ProcureFlow Complete Guide — Well-formatted Word document
covering every feature, use case, pain point, AI layer, and n8n integration.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colors ──────────────────────────────────────────────────────────
INDIGO = RGBColor(0x63, 0x66, 0xF1)
DARK = RGBColor(0x1C, 0x1C, 0x1F)
GRAY = RGBColor(0x71, 0x71, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xEF, 0x44, 0x44)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
BLUE = RGBColor(0x3B, 0x82, 0xF6)


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells, bold=False, header=False):
    """Add a row to a table with optional formatting."""
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        if bold or header:
            run.bold = True
        if header:
            set_cell_shading(cell, '6366F1')
            run.font.color.rgb = WHITE
    return row


def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a well-formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = WHITE
        set_cell_shading(cell, '6366F1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_data in rows:
        add_table_row(table, row_data)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table


def add_heading(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INDIGO if level <= 2 else DARK
    return h


def add_pain_point(doc, text):
    """Add a pain point callout."""
    p = doc.add_paragraph()
    run = p.add_run('PROBLEMA: ')
    run.bold = True
    run.font.color.rgb = RED
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK


def add_solution(doc, text):
    """Add a solution callout."""
    p = doc.add_paragraph()
    run = p.add_run('SOLUZIONE: ')
    run.bold = True
    run.font.color.rgb = GREEN
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK


def add_use_case(doc, title, steps):
    """Add a use case box."""
    p = doc.add_paragraph()
    run = p.add_run(f'USE CASE: {title}')
    run.bold = True
    run.font.color.rgb = BLUE
    run.font.size = Pt(10)

    for step in steps:
        p2 = doc.add_paragraph(step, style='List Number')
        for run in p2.runs:
            run.font.size = Pt(9.5)


def add_body(doc, text):
    """Add body text."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


# ════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT
# ════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = DARK

# Heading styles
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = INDIGO

# ── COVER PAGE ──────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ProcureFlow')
run.font.size = Pt(36)
run.font.color.rgb = INDIGO
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Guida Completa della Piattaforma')
run.font.size = Pt(18)
run.font.color.rgb = GRAY

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Hub Centralizzato di Procurement per PMI Italiane')
run.font.size = Pt(13)
run.font.color.rgb = DARK
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Versione 2.0 | {datetime.date.today().strftime("%d/%m/%Y")}')
run.font.size = Pt(10)
run.font.color.rgb = GRAY

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Stack: Next.js 14 + TypeScript + Prisma + PostgreSQL + n8n + Claude AI')
run.font.size = Pt(10)
run.font.color.rgb = GRAY

doc.add_page_break()

# ── TABLE OF CONTENTS (manual) ─────────────────────────────────────

add_heading(doc, 'Indice', 1)

toc_items = [
    '1. Panoramica — Cos\'e ProcureFlow',
    '2. Architettura del Sistema',
    '3. Sistema Modulare',
    '4. Modulo Core — Richieste di Acquisto',
    '5. Modulo Core — Gestione Fornitori',
    '6. Modulo Core — Workflow Approvazioni',
    '7. Modulo Core — Dashboard e KPI',
    '8. Modulo Core — Notifiche e Comunicazione',
    '9. Modulo Fatturazione Elettronica (SDI)',
    '10. Modulo Controllo Budget',
    '11. Modulo Gare d\'Appalto',
    '12. Modulo Magazzino e Inventario',
    '13. Modulo Analytics e ROI',
    '14. Layer AI — Panoramica Intelligenza Artificiale',
    '15. AI — Chatbot Assistente ProcureFlow',
    '16. AI — SmartFill (Suggerimenti Automatici)',
    '17. AI — Insight Cards (Intelligence Dashboard)',
    '18. AI — Forecast Inventario e Alert Riordino',
    '19. AI — OCR Fatture (Vision Parser)',
    '20. AI — Classificazione Email',
    '21. Automazione n8n — Workflow Esterni',
    '22. Sicurezza e Autenticazione',
    '23. Onboarding e Configurazione Iniziale',
    '24. Glossario Tecnico',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    for r in p.runs:
        r.font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. PANORAMICA
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Panoramica — Cos\'e ProcureFlow', 1)

add_body(doc,
    'ProcureFlow e un hub centralizzato di procurement progettato specificamente per PMI italiane. '
    'Gestisce l\'intero ciclo di vita degli acquisti: dalla richiesta d\'acquisto all\'approvazione, '
    'dall\'ordine alla consegna, dalla fatturazione elettronica alla riconciliazione contabile.')

add_body(doc,
    'Il sistema e modulare: ogni cliente attiva solo i moduli di cui ha bisogno. Il modulo Core e sempre attivo; '
    'gli altri (Fatturazione, Budget, Analytics, Gare, Magazzino) si abilitano a runtime '
    'e controllano automaticamente la navigazione, le API e le tab della dashboard.')

add_heading(doc, 'Il problema che risolve', 2)

add_pain_point(doc,
    'Nelle PMI italiane il procurement e frammentato: email, fogli Excel, telefonate, '
    'portali fornitori diversi. Ogni acquisto attraversa piu persone e strumenti senza un filo conduttore. '
    'Il risultato: ritardi, errori, spese fuori controllo e zero visibilita.')

add_solution(doc,
    'ProcureFlow centralizza tutto in un\'unica piattaforma con flussi automatizzati, '
    'approvazioni intelligenti, matching fatture e AI integrata. '
    'Dal bisogno al pagamento in un flusso tracciabile.')

add_heading(doc, 'Il ciclo completo: dal bisogno al pagamento', 2)

add_body(doc, 'ProcureFlow copre l\'intero ciclo in 10 passaggi:')

cycle_steps = [
    ['1', 'Bisogno', 'L\'utente identifica la necessita di un acquisto'],
    ['2', 'Richiesta', 'Crea la richiesta con articoli, fornitore, importo, priorita'],
    ['3', 'Budget Check', 'Il sistema verifica automaticamente la capienza del budget'],
    ['4', 'Approvazione', 'Workflow automatico basato sull\'importo (auto/manager/direttore)'],
    ['5', 'Ordine', 'L\'ordine viene tracciato con riferimento fornitore e tracking'],
    ['6', 'Consegna', 'Monitoraggio automatico con alert per ritardi'],
    ['7', 'Fattura SDI', 'Ricezione automatica della fattura elettronica dal SDI'],
    ['8', 'Matching', 'Abbinamento automatico fattura-ordine (5 livelli di confidenza)'],
    ['9', 'Riconciliazione', 'Three-way matching: ordine vs fattura vs merce ricevuta'],
    ['10', 'Chiusura', 'Approvazione pagamento, budget aggiornato, tutto tracciato'],
]
create_styled_table(doc, ['#', 'Fase', 'Descrizione'], cycle_steps, [1.5, 3, 12])

add_heading(doc, 'Per chi e ProcureFlow', 2)

create_styled_table(doc, ['Ruolo', 'Beneficio principale'], [
    ['Responsabile Acquisti', 'Visibilita completa sul ciclo, dashboard con KPI, zero email perse'],
    ['Manager/Direttore', 'Approvazioni rapide da qualsiasi dispositivo, controllo budget in tempo reale'],
    ['Richiedente', 'Crea richieste in 2 minuti, tracking automatico, notifiche su ogni cambio stato'],
    ['Amministrazione', 'Fatture abbinate automaticamente, riconciliazione 3-way, compliance SDI'],
    ['IT/Ops', 'Deploy modulare, API RESTful, integrazione n8n, zero manutenzione'],
], [4, 12])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 2. ARCHITETTURA
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Architettura del Sistema', 1)

add_body(doc, 'ProcureFlow utilizza un\'architettura moderna a 4 livelli:')

create_styled_table(doc, ['Livello', 'Tecnologia', 'Responsabilita'], [
    ['Frontend', 'React + TanStack Query + Next.js 14', 'UI interattiva, caching client-side, routing'],
    ['API Layer', 'Next.js API Routes + Zod', 'Validazione input, auth JWT, module guard'],
    ['Service Layer', 'TypeScript (business logic pura)', 'Logica di dominio, transazioni, calcoli'],
    ['Database', 'PostgreSQL (Supabase) + Prisma ORM', 'Persistenza, indici ottimizzati, PgBouncer'],
    ['Automazione', 'n8n (self-hosted)', 'Email ingestion, approval workflow, delivery monitoring'],
    ['AI Layer', 'Anthropic Claude API', 'Chatbot, SmartFill, insights, forecast, OCR fatture'],
], [3, 5.5, 8])

add_heading(doc, 'Flusso dati', 2)

add_body(doc,
    'Il browser invia richieste HTTP alle API Routes di Next.js. '
    'Le API validano l\'input con Zod, verificano autenticazione e ruolo, '
    'poi delegano al Service Layer. I servizi eseguono la logica di business '
    'e interagiscono con il database via Prisma ORM. '
    'n8n comunica col sistema tramite webhook HTTP autenticati con HMAC.')

add_heading(doc, 'Stack tecnologico completo', 2)

create_styled_table(doc, ['Componente', 'Tecnologia', 'Perche questa scelta'], [
    ['Framework', 'Next.js 14 (App Router)', 'SSR + CSR ibrido, API integrate, eccellente DX'],
    ['Linguaggio', 'TypeScript (strict mode)', 'Type safety end-to-end, refactoring sicuro'],
    ['Styling', 'Tailwind CSS', 'Design system coerente, zero CSS custom da mantenere'],
    ['Database', 'PostgreSQL su Supabase', 'Affidabile, scalabile, RLS, backup automatici'],
    ['ORM', 'Prisma', 'Schema dichiarativo, migrazioni, type-safe queries'],
    ['Auth', 'NextAuth.js (JWT)', 'RBAC 4 ruoli, MFA, session management'],
    ['State', 'TanStack Query', 'Cache client, optimistic updates, sync automatico'],
    ['Automazione', 'n8n (self-hosted)', 'Visual workflow builder, 400+ integrazioni'],
    ['AI', 'Anthropic Claude API', 'Multimodale (testo + immagini), italiano nativo'],
    ['Monitoring', 'Sentry', 'Error tracking, performance monitoring'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 3. SISTEMA MODULARE
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Sistema Modulare', 1)

add_body(doc,
    'ProcureFlow e costruito attorno a un sistema modulare: ogni cliente attiva solo i moduli necessari '
    'tramite una configurazione in database (DeployConfig). Il modulo Core e sempre attivo.')

add_pain_point(doc,
    'Le piattaforme enterprise offrono tutto o niente: interfacce sovraccariche di funzionalita '
    'non necessarie, costi elevati, tempi di adozione lunghi.')

add_solution(doc,
    'Con ProcureFlow attivi solo cio che ti serve. L\'interfaccia resta pulita, '
    'i costi sono proporzionali, l\'adozione e graduale.')

create_styled_table(doc, ['Modulo', 'ID', 'Cosa controlla', 'Default'], [
    ['Core', 'core', 'Richieste, Fornitori, Approvazioni, Utenti, Dashboard', 'Sempre attivo'],
    ['Fatturazione', 'invoicing', 'Fatture SDI, matching, riconciliazione 3-way', 'Opzionale'],
    ['Budget', 'budgets', 'Plafond per centro di costo, enforcement spesa', 'Opzionale'],
    ['Analytics', 'analytics', 'Dashboard ROI, trend spesa, report', 'Opzionale'],
    ['Gare', 'tenders', 'Gare d\'appalto, Go/No-Go, compliance CIG/CUP', 'Opzionale'],
    ['Magazzino', 'inventory', 'Materiali, lotti, movimenti, inventari fisici', 'Opzionale'],
], [2.5, 2, 8, 3])

add_body(doc,
    'Quando un modulo e disabilitato: la voce di navigazione scompare dalla sidebar, '
    'le API restituiscono 404, la tab corrispondente nella dashboard non viene renderizzata, '
    'nessun dato viene fetchato. L\'utente non vede mai funzionalita che non ha acquistato.')

add_use_case(doc, 'Attivazione modulo Budget', [
    'L\'admin va in Impostazioni > Moduli',
    'Clicca su "Attiva" accanto al modulo Budget',
    'La sidebar mostra immediatamente la voce "Budget"',
    'La dashboard aggiunge la tab "Budget" con allocato/speso/impegnato',
    'Le API /api/budgets diventano accessibili',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 4. RICHIESTE DI ACQUISTO
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Modulo Core — Richieste di Acquisto', 1)

add_body(doc,
    'Il cuore del sistema. Ogni acquisto nasce come PurchaseRequest (PR) con un ciclo di vita '
    'a 12 stati, dal DRAFT alla CLOSED. Ogni azione e tracciata nella timeline.')

add_pain_point(doc,
    'Le richieste di acquisto viaggiano via email: si perdono, non hanno un codice univoco, '
    'nessuno sa a che punto e un ordine, i ritardi si accumulano senza visibilita.')

add_solution(doc,
    'Ogni richiesta ha un codice univoco (PR-2026-NNNNN), una timeline visuale completa, '
    'notifiche automatiche ad ogni cambio stato, e uno storico immutabile di tutte le azioni.')

add_heading(doc, 'Cosa puoi fare', 2)

bullets = [
    ('Creare richieste', 'con titolo, descrizione, articoli dettagliati (quantita, prezzo, SKU), fornitore, priorita, centro di costo, data necessita'),
    ('Tracciare ogni passaggio', 'dalla bozza alla chiusura con timeline visuale completa'),
    ('Allegare documenti', 'PDF, DOCX, XLSX, immagini fino a 10 MB'),
    ('Commentare', 'internamente o con il fornitore direttamente sulla richiesta'),
    ('Compliance italiana', 'campi CIG, CUP, flag MePA, numero ODA supportati nativamente'),
]
for prefix, text in bullets:
    add_bullet(doc, text, prefix)

add_heading(doc, 'Macchina a stati', 2)

create_styled_table(doc, ['Stato', 'Significato', 'Transizione successiva'], [
    ['DRAFT', 'Bozza in lavorazione', 'SUBMITTED'],
    ['SUBMITTED', 'Inviata per approvazione', 'PENDING_APPROVAL'],
    ['PENDING_APPROVAL', 'In attesa di decisione', 'APPROVED / REJECTED'],
    ['APPROVED', 'Approvata, pronta per ordine', 'ORDERED'],
    ['REJECTED', 'Rifiutata con motivazione', '(terminale)'],
    ['ORDERED', 'Ordine inviato al fornitore', 'SHIPPED'],
    ['SHIPPED', 'Merce spedita', 'DELIVERED'],
    ['DELIVERED', 'Merce ricevuta', 'INVOICED'],
    ['INVOICED', 'Fattura ricevuta e abbinata', 'RECONCILED'],
    ['RECONCILED', 'Riconciliazione completata', 'CLOSED'],
    ['CLOSED', 'Ciclo completato', '(terminale)'],
    ['CANCELLED', 'Annullata', '(terminale)'],
    ['ON_HOLD', 'Sospesa temporaneamente', '(qualsiasi stato precedente)'],
])

add_use_case(doc, 'Test — Creazione e approvazione richiesta', [
    'Login come mario.rossi@saiflow.com (ruolo REQUESTER)',
    'Vai su Richieste > Nuova Richiesta',
    'Compila: Titolo "10 Sedie ergonomiche", Fornitore "Officine Meccaniche SRL", Priorita MEDIUM',
    'Aggiungi articolo: "Sedia ergonomica X200", qty 10, prezzo 245.00 EUR',
    'Il sistema calcola automaticamente il totale: 2.450,00 EUR',
    'Clicca "Invia" — stato passa a SUBMITTED > PENDING_APPROVAL',
    'Login come laura.manager@saiflow.com (MANAGER)',
    'Vai su Approvazioni — la richiesta appare nella coda',
    'Clicca "Approva" con nota "OK per Q2" — stato passa a APPROVED',
    'Mario riceve notifica "La tua richiesta PR-2026-00043 e stata approvata"',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 5. GESTIONE FORNITORI
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Modulo Core — Gestione Fornitori', 1)

add_pain_point(doc,
    'Le informazioni sui fornitori sono sparse: rubrica personale, Excel condiviso, '
    'email con contatti diversi. Quando un dipendente deve ordinare non sa chi contattare.')

add_solution(doc,
    'Anagrafica centralizzata con ragione sociale, P.IVA (per matching fatture automatico), '
    'contatti multipli, rating interno, categorie merceologiche e storico completo.')

create_styled_table(doc, ['Campo', 'Descrizione', 'Esempio'], [
    ['Codice', 'Identificativo univoco interno', 'FIM-001'],
    ['Ragione sociale', 'Nome legale del fornitore', 'TechParts Italia SRL'],
    ['P.IVA', 'Per matching automatico fatture SDI', 'IT12345678901'],
    ['Categoria', 'Classificazione merceologica (array)', 'Hardware, Networking'],
    ['Tipo portale', 'Modalita di interazione', 'API / Email / Marketplace / Telefono'],
    ['Rating', 'Valutazione interna 0-5', '4.2'],
    ['Termini pagamento', 'Condizioni contrattuali', '30gg DFFM'],
    ['Stato', 'Stato operativo', 'ACTIVE / INACTIVE / BLACKLISTED'],
])

add_use_case(doc, 'Test — Creare e cercare un fornitore', [
    'Login come admin — vai su Fornitori > Nuovo Fornitore',
    'Compila: Nome "Grafica Express", P.IVA "IT98765432100", Categoria "Stampa"',
    'Salva — il sistema assegna il codice VEN-00051',
    'Torna alla lista — cerca "Grafica" nella barra di ricerca',
    'Il fornitore appare con tutti i dati e il rating "Non valutato"',
    'Clicca sul fornitore — vedi la sezione "Richieste" (vuota per ora) e "Fatture"',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 6. WORKFLOW APPROVAZIONI
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '6. Modulo Core — Workflow Approvazioni', 1)

add_pain_point(doc,
    'Le approvazioni via email sono lente: il manager e fuori ufficio, l\'email finisce in spam, '
    'nessuno sa chi deve approvare cosa. Le richieste restano bloccate per giorni.')

add_solution(doc,
    'Approvazioni automatiche basate su soglie di importo. Le piccole spese passano '
    'senza intervento umano, le altre generano notifiche immediate agli approvatori giusti.')

create_styled_table(doc, ['Importo', 'Tipo approvazione', 'Tempo tipico'], [
    ['< 500 EUR', 'Automatica — nessun intervento umano', 'Istantaneo'],
    ['500 - 4.999 EUR', 'Approvazione Manager', '< 1 ora con notifica'],
    ['>= 5.000 EUR', 'Approvazione Direttore', '< 4 ore con escalation'],
])

add_heading(doc, 'Funzionalita', 2)
add_bullet(doc, 'in attesa con filtri per stato, importo e priorita', 'Coda approvazioni')
add_bullet(doc, 'con un click, aggiungendo note motivazionali', 'Approva o rifiuta')
add_bullet(doc, 'immediate quando una richiesta richiede la tua decisione', 'Notifiche')
add_bullet(doc, 'l\'approvazione a un collega in caso di assenza', 'Delega')
add_bullet(doc, 'di tutte le decisioni prese con data e motivazione', 'Storico completo')

add_use_case(doc, 'Test — Flusso approvazione completo', [
    'Crea una richiesta da 200 EUR — viene auto-approvata (verifica stato APPROVED)',
    'Crea una richiesta da 1.500 EUR — rimane PENDING_APPROVAL',
    'Login come Manager — trovi la richiesta nella coda approvazioni',
    'Approva con nota "OK budget Q2" — stato diventa APPROVED',
    'Il richiedente riceve notifica e vede l\'evento nella timeline',
    'Crea una richiesta da 8.000 EUR — richiede approvazione Direttore',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 7. DASHBOARD E KPI
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '7. Modulo Core — Dashboard e KPI', 1)

add_pain_point(doc,
    'Per capire lo stato degli acquisti bisogna aprire email, Excel, ERP. '
    'Nessuna visione d\'insieme: le informazioni critiche sono sparse ovunque.')

add_solution(doc,
    'Dashboard interattiva con KPI in tempo reale, tab modulari per ogni area, '
    'grafici trend, insight AI e accesso rapido alle azioni pendenti.')

add_heading(doc, 'I 4 KPI principali', 2)

create_styled_table(doc, ['KPI', 'Cosa mostra', 'Perche conta'], [
    ['Richieste Attive', 'Numero in corso + trend vs mese precedente', 'Carico di lavoro attuale'],
    ['In Attesa di Approvazione', 'Richieste che aspettano decisione', 'Colli di bottiglia da sbloccare'],
    ['Spesa Mensile', 'Totale speso + % rispetto al budget', 'Controllo finanziario'],
    ['Consegne in Ritardo', 'Ordini con data consegna superata', 'Problemi da gestire subito'],
])

add_heading(doc, 'Tab modulari della dashboard', 2)

create_styled_table(doc, ['Tab', 'Contenuto', 'Modulo richiesto'], [
    ['Panoramica', 'Richieste recenti, prossime consegne, trend spesa, distribuzione stati', 'Core (sempre)'],
    ['Fatture', 'Non abbinate, aging (0-30/30-60/60-90/90+ gg), ordinato vs fatturato', 'Fatturazione'],
    ['Budget', 'Allocato/speso/impegnato per centro di costo, burn rate', 'Budget'],
    ['Gare', 'Pipeline attiva, valore totale, scadenze, tasso di vittoria', 'Gare'],
    ['Magazzino', 'Valore stock, alert scorte basse, trend movimenti', 'Magazzino'],
    ['Analisi', 'Top 5 fornitori per spesa, trend richieste 6 mesi, mini ROI', 'Analytics'],
])

add_use_case(doc, 'Test — Verifica dashboard completa', [
    'Login come admin — la dashboard mostra tutti i KPI',
    'Verifica che il contatore "Richieste Attive" corrisponda alla lista reale',
    'Clicca sulla tab "Analisi" — vedrai trend spesa per fornitore',
    'Verifica che le Insight AI card mostrino suggerimenti (se ci sono dati)',
    'Disabilita un modulo da Impostazioni — la tab corrispondente scompare dalla dashboard',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 8. NOTIFICHE
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '8. Modulo Core — Notifiche e Comunicazione', 1)

add_pain_point(doc,
    'Le email importanti si perdono nella casella piena. '
    'I cambi di stato sugli ordini non vengono comunicati. '
    'I ritardi si scoprono troppo tardi.')

add_solution(doc,
    'Sistema di notifiche in-app con 10 tipi diversi, badge visuale nell\'header, '
    'deep link diretto alla risorsa e segna come letto con un click.')

create_styled_table(doc, ['Tipo', 'Quando scatta', 'Destinatario'], [
    ['Approvazione richiesta', 'Nuova richiesta da approvare', 'Manager/Direttore'],
    ['Richiesta approvata', 'La richiesta e stata approvata', 'Richiedente'],
    ['Richiesta rifiutata', 'La richiesta e stata rifiutata', 'Richiedente'],
    ['Consegna in ritardo', 'Data consegna superata', 'Richiedente + Manager'],
    ['Promemoria consegna', 'Consegna prevista entro 3 giorni', 'Richiedente'],
    ['Cambio stato', 'Stato PR cambiato', 'Richiedente'],
    ['Nuovo commento', 'Commento aggiunto alla richiesta', 'Partecipanti'],
    ['Fattura ricevuta', 'Nuova fattura dal fornitore', 'Amministrazione'],
    ['Fattura abbinata', 'Fattura collegata a un ordine', 'Richiedente'],
    ['Alert budget', 'Soglia di spesa raggiunta', 'Manager + Admin'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 9. FATTURAZIONE
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '9. Modulo Fatturazione Elettronica (SDI)', 1)

add_pain_point(doc,
    'Le fatture arrivano via PEC, vanno scaricate manualmente, registrate in Excel, '
    'poi confrontate a mano con l\'ordine. Il rischio di errori e altissimo, '
    'i tempi di pagamento si allungano, le discrepanze si scoprono tardi.')

add_solution(doc,
    'Ricezione automatica dal SDI, parsing completo del formato FatturaPA, '
    'matching a 5 livelli con le richieste di acquisto, riconciliazione 3-way automatica. '
    'Il 70-80% delle fatture viene gestito senza intervento umano.')

add_heading(doc, 'Il flusso completo', 2)
add_body(doc, 'Fattura XML dal SDI > Parsing automatico > Abbinamento all\'ordine > Riconciliazione 3-way > Pagamento')

add_heading(doc, 'Algoritmo di matching a 5 livelli', 2)

create_styled_table(doc, ['Livello', 'Criterio', 'Confidenza', 'Azione'], [
    ['1', 'Codice PR nella causale (es. "PR-2026-00042")', '95-100%', 'Auto-match'],
    ['2', 'Codice PR nei riferimenti DatiOrdineAcquisto', '85%', 'Auto-match'],
    ['3', 'P.IVA fornitore + importo simile (+-10%)', '60-70%', 'Suggerito'],
    ['4', 'P.IVA fornitore + finestra 90 giorni', '40-50%', 'Suggerito'],
    ['5', 'Nessun match trovato', '<40%', 'Manuale'],
])

add_body(doc, 'Soglia auto-match: confidenza >= 80%. Sotto questa soglia l\'operatore conferma.')

add_heading(doc, 'Three-Way Matching (Riconciliazione a 3 vie)', 2)

add_body(doc, 'Il sistema confronta tre fonti: quanto ordinato (PR approvata), quanto ricevuto (magazzino), quanto fatturato.')

create_styled_table(doc, ['Discrepanza', 'Soglia', 'Azione'], [
    ['< 2%', 'Sotto tolleranza', 'Approvazione automatica'],
    ['2% - 5%', 'Warning', 'Conferma manuale richiesta'],
    ['> 5%', 'Critica', 'Review obbligatoria, blocco pagamento'],
])

add_heading(doc, 'Upload manuale con OCR AI', 2)
add_body(doc,
    'Per fatture non XML (PDF, immagini), il sistema usa Claude Vision per estrarre automaticamente: '
    'numero fattura, data, importi, righe, P.IVA. L\'AI restituisce un punteggio di confidenza '
    'e l\'operatore puo correggere prima della conferma. Vedi sezione 19 per dettagli.')

add_use_case(doc, 'Test — Ricezione e riconciliazione fattura', [
    'Crea e approva una richiesta di acquisto per "TechParts Italia" da 2.945 EUR',
    'Segna la richiesta come ORDERED, poi DELIVERED',
    'Simula l\'arrivo di una fattura SDI tramite webhook /api/webhooks/sdi-invoice',
    'La fattura viene parsata automaticamente, il codice PR viene trovato nella causale',
    'Verifica in Fatture che lo stato sia "MATCHED" con confidenza 95%',
    'La riconciliazione confronta: ordinato 2.945, fatturato 2.945, ricevuti tutti — APPROVED',
    'L\'amministrazione vede la fattura pronta per il pagamento senza intervento manuale',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 10. BUDGET
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '10. Modulo Controllo Budget', 1)

add_pain_point(doc,
    'I budget esistono su fogli Excel aggiornati manualmente. Nessuno sa in tempo reale '
    'quanto rimane. Gli sforamenti si scoprono a fine mese quando e troppo tardi.')

add_solution(doc,
    'Budget per centro di costo con controllo automatico alla sottomissione, '
    'due modalita di enforcement (SOFT/HARD), burn rate e previsione esaurimento in tempo reale.')

add_heading(doc, 'Come funziona il calcolo', 2)

create_styled_table(doc, ['Componente', 'Definizione', 'Esempio'], [
    ['Allocato', 'Budget totale assegnato al periodo', '50.000 EUR'],
    ['Speso', 'Richieste ordinate/consegnate/fatturate', '28.000 EUR'],
    ['Impegnato', 'Richieste inviate/approvate ma non ordinate', '8.500 EUR'],
    ['Disponibile', 'Allocato - Speso - Impegnato', '13.500 EUR'],
    ['Burn Rate', 'Velocita di consumo giornaliera', '1.200 EUR/giorno'],
    ['Esaurimento', 'Data prevista esaurimento budget', '15 aprile 2026'],
])

add_heading(doc, 'Modalita di enforcement', 2)
add_bullet(doc, 'Mostra un avviso arancione ma permette di procedere con la richiesta', 'SOFT:')
add_bullet(doc, 'Blocca completamente la sottomissione se il budget e insufficiente', 'HARD:')

add_use_case(doc, 'Test — Controllo budget attivo', [
    'Vai in Budget > crea un budget per centro di costo "IT" da 10.000 EUR, modalita HARD',
    'Crea una richiesta da 8.000 EUR per il centro di costo "IT" — viene accettata (2.000 disponibili)',
    'Crea una seconda richiesta da 3.000 EUR per "IT" — il sistema BLOCCA la sottomissione',
    'Appare l\'errore: "Budget insufficiente: disponibile 2.000 EUR, richiesto 3.000 EUR"',
    'Cambia la modalita del budget a SOFT — ora la seconda richiesta passa con un warning',
    'Nella dashboard tab Budget, verifica il burn rate e la data prevista di esaurimento',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 11. GARE D'APPALTO
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '11. Modulo Gare d\'Appalto', 1)

add_pain_point(doc,
    'La gestione delle gare e su cartelle condivise e spreadsheet. '
    'Scadenze mancate, documenti persi, nessun processo strutturato per il Go/No-Go.')

add_solution(doc,
    'Pipeline visuale con 8 stati, processo Go/No-Go strutturato, '
    'gestione documenti integrata, compliance CIG/CUP, timeline completa.')

create_styled_table(doc, ['Stato', 'Significato'], [
    ['SCOUTING', 'Gara individuata, in fase di valutazione'],
    ['ANALYSIS', 'Analisi del bando, requisiti tecnici e economici'],
    ['GO_NO_GO', 'Decisione strutturata se partecipare o meno'],
    ['PREPARATION', 'Preparazione offerta tecnica ed economica'],
    ['SUBMITTED', 'Offerta inviata alla stazione appaltante'],
    ['EVALUATION', 'In valutazione dalla commissione'],
    ['WON', 'Gara aggiudicata'],
    ['LOST', 'Gara non aggiudicata'],
])

add_use_case(doc, 'Test — Ciclo gara completo', [
    'Vai in Gare > Nuova Gara',
    'Compila: Titolo "Fornitura Cavi Rete PA Lombardia", CIG "8521734AB1", Importo base 45.000 EUR',
    'La gara inizia in stato SCOUTING',
    'Avanza ad ANALYSIS — aggiungi note sui requisiti tecnici',
    'Avanza a GO_NO_GO — documenta la decisione ("Margine 22%, competenze OK, GO")',
    'Prepara l\'offerta (PREPARATION), allega i documenti, poi SUBMITTED',
    'Simula l\'esito: segna come WON — la gara appare nella pipeline come vinta',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 12. MAGAZZINO
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '12. Modulo Magazzino e Inventario', 1)

add_pain_point(doc,
    'Le scorte sono gestite a occhio o su Excel. Non si sa cosa c\'e in magazzino, '
    'i materiali finiscono senza preavviso, gli inventari fisici sono un incubo.')

add_solution(doc,
    'Gestione completa di materiali, magazzini, zone, lotti con tracciabilita, '
    'movimenti in/out, riserve, inventari fisici con conteggio e rettifiche, '
    'e forecast AI per previsione consumo (vedi sezione 18).')

add_heading(doc, 'Entita gestite', 2)

create_styled_table(doc, ['Entita', 'Cosa rappresenta', 'Esempio'], [
    ['Materiale', 'Articolo a stock con codice, UM, costo, scorte min/max', 'Cavo RG59 (MAT-CAV-00001)'],
    ['Magazzino', 'Sede fisica di stoccaggio', 'Magazzino Centrale Milano'],
    ['Zona', 'Suddivisione logica del magazzino', 'Scaffale A / Corsia 3'],
    ['Lotto', 'Stock specifico con tracciabilita (scadenza, costo)', 'Lotto #2026-001, 500 pezzi'],
    ['Movimento', 'Entrata/uscita/trasferimento con causale', 'OUTBOUND - Prelievo per PR-2026-00042'],
    ['Riserva', 'Quantita bloccata per una specifica richiesta/gara', '100 pz riservati per gara CIG-001'],
    ['Inventario', 'Conteggio fisico periodico con rettifiche', 'Inventario semestrale giugno 2026'],
])

add_heading(doc, 'Livelli di stock', 2)

create_styled_table(doc, ['Stato', 'Condizione', 'Indicatore'], [
    ['OK', 'Stock > scorta minima * 1.5', 'Verde'],
    ['LOW', 'Stock tra scorta minima e scorta minima * 1.5', 'Arancione'],
    ['CRITICAL', 'Stock <= scorta minima', 'Rosso'],
    ['OUT', 'Stock = 0', 'Rosso con icona allarme'],
])

add_use_case(doc, 'Test — Ciclo magazzino completo', [
    'Vai in Magazzino > Materiali — vedi la lista con badge colorato stock level',
    'Clicca su "Cavo RG59" — tab Generale mostra codice, UM, costi, scorte',
    'Tab Giacenze — vedi stock per magazzino',
    'Tab Lotti — lotti attivi con quantita e scadenza',
    'Tab Movimenti — storico entrate/uscite con causale',
    'Tab Forecast AI — proiezione consumo 3 mesi, giorni rimanenti, bottone Analisi AI',
    'Clicca "Analisi AI" — Claude analizza lo storico e suggerisce quantita di riordino',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 13. ANALYTICS E ROI
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '13. Modulo Analytics e ROI', 1)

add_pain_point(doc,
    'Il management non ha visibilita sul ritorno dell\'investimento del procurement. '
    'I report sono manuali, i dati sono vecchi di settimane, le decisioni sono a intuito.')

add_solution(doc,
    'Dashboard analytics con metriche ROI calcolate in tempo reale: tempo risparmiato, '
    'efficienza approvazioni, compliance budget, savings negoziazione. '
    'Ogni metrica confronta il prima (manuale) e il dopo (ProcureFlow).')

add_heading(doc, 'Metriche principali', 2)

create_styled_table(doc, ['Metrica', 'Cosa misura', 'Benchmark manuale'], [
    ['Ciclo medio richiesta', 'Giorni da SUBMITTED a DELIVERED', '14 giorni'],
    ['Tempo medio approvazione', 'Ore da PENDING a decisione', '48 ore'],
    ['Risparmio negoziazione', '% differenza stimato vs effettivo', '0% (no tracking)'],
    ['Compliance budget', '% richieste con budget disponibile', 'Sconosciuta'],
    ['Tasso auto-match fatture', '% fatture abbinate automaticamente', '0% (tutto manuale)'],
    ['Consegne puntuali', '% ordini consegnati entro la data prevista', '65%'],
    ['Ore risparmiate', 'Ore umane risparmiate dall\'automazione', 'N/A'],
])

add_heading(doc, 'Export CSV', 2)
add_body(doc,
    'Ogni metrica e esportabile in CSV con un click. Il file contiene le metriche attuali, '
    'i benchmark manuali e il delta di miglioramento.')

add_use_case(doc, 'Test — Consultare analytics', [
    'Vai su Analytics dalla sidebar',
    'Verifica i 6 summary card: ciclo medio, approvazione, savings, compliance, auto-match, consegne',
    'I grafici mostrano trend mensile per ogni metrica',
    'Clicca "Esporta CSV" — scarica il file con tutte le metriche e benchmark',
    'Verifica che le ore risparmiate includano tutte le fonti (richieste, email, fatture)',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 14. LAYER AI — PANORAMICA
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '14. Layer AI — Panoramica Intelligenza Artificiale', 1)

add_body(doc,
    'ProcureFlow integra l\'intelligenza artificiale di Anthropic Claude in 6 aree funzionali. '
    'Ogni funzione AI e progettata per risolvere un problema specifico del procurement, '
    'non come feature decorativa ma come acceleratore operativo concreto.')

create_styled_table(doc, ['Funzione AI', 'Modello', 'Dove agisce', 'Pain point risolto'], [
    ['Chatbot Assistente', 'Claude Haiku 4.5', 'Widget chat nella dashboard', 'Cercare informazioni tra centinaia di richieste/fornitori'],
    ['SmartFill', 'Claude Sonnet 4.5', 'Form nuova richiesta', 'Compilare campi ripetitivi per acquisti ricorrenti'],
    ['Insight Cards', 'Claude Sonnet 4.5', 'Dashboard principale', 'Identificare anomalie e opportunita nei dati di spesa'],
    ['Forecast Inventario', 'Claude Sonnet 4.5', 'Dettaglio materiale', 'Prevedere quando un materiale finira e quanto riordinare'],
    ['OCR Fatture', 'Claude Vision', 'Upload fattura manuale', 'Digitalizzare fatture cartacee/PDF non strutturati'],
    ['Classificazione Email', 'Claude Sonnet 4.5', 'n8n email ingestion', 'Capire se un\'email contiene info su un ordine'],
])

add_heading(doc, 'Architettura AI', 2)

add_body(doc,
    'Tutte le chiamate AI passano attraverso un client centralizzato (claude-client.ts) '
    'che gestisce: selezione del modello, rate limiting, error handling, logging. '
    'I prompt di sistema sono definiti in un file dedicato (prompts.ts) per mantenibilita. '
    'L\'API key Anthropic e configurata in .env.local (ANTHROPIC_API_KEY).')

add_heading(doc, 'Scelta dei modelli', 2)

create_styled_table(doc, ['Modello', 'Uso', 'Perche'], [
    ['Haiku 4.5', 'Chatbot', '3x piu veloce, perfetto per conversazione in tempo reale'],
    ['Sonnet 4.5', 'SmartFill, Insights, Forecast, Email', 'Miglior rapporto qualita/costo per task analitici'],
    ['Claude Vision', 'OCR Fatture', 'Capacita multimodale per estrarre dati da immagini/PDF'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 15. CHATBOT
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '15. AI — Chatbot Assistente ProcureFlow', 1)

add_pain_point(doc,
    'Per trovare informazioni bisogna navigare tra pagine, applicare filtri, cercare manualmente. '
    '"Quanto ho speso con TechParts quest\'anno?" richiede export Excel + formule.')

add_solution(doc,
    'Chatbot conversazionale con accesso a tutti i dati della piattaforma. '
    'Puoi chiedere in linguaggio naturale e ricevere risposte istantanee con dati reali.')

add_heading(doc, 'Come funziona', 2)

add_body(doc,
    'Il chatbot usa un sistema di tool-calling: Claude riceve la domanda dell\'utente '
    'e decide autonomamente quali "strumenti" usare per rispondere. '
    'Gli strumenti disponibili sono:')

create_styled_table(doc, ['Tool', 'Cosa fa', 'Esempio domanda'], [
    ['search_requests', 'Cerca richieste per stato, fornitore, importo', '"Mostra le richieste pendenti"'],
    ['get_request_detail', 'Dettaglio completo di una PR specifica', '"Dettagli su PR-2026-00042"'],
    ['search_vendors', 'Cerca fornitori per nome, categoria, stato', '"Fornitori di hardware attivi"'],
    ['get_dashboard_stats', 'KPI della dashboard in tempo reale', '"Come sta andando il procurement?"'],
    ['search_invoices', 'Cerca fatture per stato, fornitore, importo', '"Fatture non pagate del mese"'],
    ['create_request', 'Crea una nuova richiesta di acquisto', '"Crea una richiesta per 10 sedie"'],
])

add_heading(doc, 'Sicurezza (RBAC)', 2)
add_body(doc,
    'Ogni tool rispetta i permessi dell\'utente: un REQUESTER non puo accedere ai dati '
    'di tutti gli utenti, un VIEWER non puo creare richieste. Il tool create_request '
    'richiede conferma esplicita prima di eseguire l\'azione.')

add_heading(doc, 'Streaming SSE', 2)
add_body(doc,
    'Le risposte vengono inviate in streaming (Server-Sent Events) per un\'esperienza '
    'fluida. Il chatbot inizia a rispondere immediatamente, senza attese.')

add_use_case(doc, 'Test — Interazione chatbot', [
    'Apri il widget chat (icona in basso a destra)',
    'Scrivi: "Quanto ho speso con TechParts quest\'anno?"',
    'Il chatbot cerca le fatture/richieste e risponde con importo totale e breakdown',
    'Scrivi: "Crea una richiesta per 10 sedie ergonomiche a 245 EUR l\'una, fornitore Officine Meccaniche"',
    'Il chatbot mostra un riepilogo e chiede conferma prima di creare',
    'Conferma — la richiesta viene creata e il chatbot fornisce il codice PR',
    'Scrivi: "Quante richieste sono in attesa di approvazione?" — risposta immediata con conteggio',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 16. SMARTFILL
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '16. AI — SmartFill (Suggerimenti Automatici)', 1)

add_pain_point(doc,
    'Ogni richiesta di acquisto richiede la compilazione di 10+ campi. '
    'Per acquisti ricorrenti (cancelleria, IT) si riscrivono sempre le stesse cose.')

add_solution(doc,
    'SmartFill analizza lo storico delle richieste precedenti e suggerisce automaticamente '
    'fornitore, articoli, quantita, prezzi e centro di costo basandosi sul titolo inserito.')

add_heading(doc, 'Come funziona', 2)
add_body(doc,
    'Quando l\'utente inserisce il titolo della richiesta (es. "Toner stampante HP"), '
    'il pannello SmartFill mostra suggerimenti basati su richieste simili nel passato. '
    'L\'utente accetta il suggerimento con un click e i campi vengono pre-compilati.')

create_styled_table(doc, ['Campo suggerito', 'Fonte', 'Esempio'], [
    ['Fornitore', 'Fornitore piu usato per articoli simili', 'Office Depot'],
    ['Articoli', 'Lista articoli da richieste simili precedenti', 'Toner HP 305A, qty 3'],
    ['Prezzo', 'Ultimo prezzo pagato + trend', '45.00 EUR'],
    ['Centro di costo', 'Piu usato per questa categoria', 'IT-STAMPANTI'],
    ['Priorita', 'Basata su frequenza e urgenza storica', 'MEDIUM'],
])

add_use_case(doc, 'Test — SmartFill in azione', [
    'Vai su Richieste > Nuova Richiesta',
    'Digita nel titolo: "Toner stampante"',
    'Sulla destra appare il pannello "Suggerimenti AI" con richieste simili passate',
    'Clicca "Applica suggerimento" — fornitore, articoli e prezzo vengono pre-compilati',
    'Modifica la quantita se necessario e invia',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 17. INSIGHT CARDS
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '17. AI — Insight Cards (Intelligence Dashboard)', 1)

add_pain_point(doc,
    'I problemi di procurement si scoprono troppo tardi: uno sforamento budget, '
    'un fornitore inaffidabile, un trend di spesa anomalo. Manca l\'analisi proattiva.')

add_solution(doc,
    'Le Insight Cards analizzano automaticamente i dati e generano avvisi proattivi '
    'direttamente nella dashboard. Ogni card ha severita, azione suggerita e deep link.')

add_heading(doc, 'Tipi di insight', 2)

create_styled_table(doc, ['Tipo', 'Icona', 'Esempio', 'Severita tipica'], [
    ['SPEND_ANOMALY', 'TrendingUp', 'Spesa con TechParts +45% vs media', 'HIGH'],
    ['VENDOR_RISK', 'ShieldAlert', 'Fornitore ABC: 3 consegne in ritardo su 5', 'CRITICAL'],
    ['SAVINGS', 'PiggyBank', 'Possibile risparmio 1.200 EUR consolidando ordini cancelleria', 'MEDIUM'],
    ['BOTTLENECK', 'Timer', 'Approvazioni bloccate: 5 richieste in attesa da >48h', 'HIGH'],
    ['BUDGET_ALERT', 'AlertTriangle', 'Budget IT al 92% — esaurimento previsto in 8 giorni', 'CRITICAL'],
])

add_heading(doc, 'Comportamento', 2)
add_bullet(doc, 'Le card appaiono nella dashboard solo quando ci sono insight attivi', 'Visibilita:')
add_bullet(doc, 'L\'utente puo nascondere una card con la X — non riappare', 'Dismissione:')
add_bullet(doc, 'Ogni card ha un link diretto all\'area pertinente', 'Azione:')
add_bullet(doc, 'CRITICAL (rosso), HIGH (arancione), MEDIUM (blu), LOW (grigio)', 'Severita:')

add_use_case(doc, 'Test — Insight cards in dashboard', [
    'Login come admin — apri la dashboard',
    'Se ci sono dati sufficienti, le Insight Cards appaiono sotto i KPI',
    'Verifica: ogni card mostra titolo, descrizione, severita (colore bordo sinistro) e icona',
    'Clicca il link "Vai alle richieste >" su una card — ti porta alla pagina corretta',
    'Clicca la X su una card — scompare e non riappare',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 18. FORECAST
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '18. AI — Forecast Inventario e Alert Riordino', 1)

add_pain_point(doc,
    'Il magazziniere scopre che un materiale e finito quando qualcuno lo cerca e non lo trova. '
    'Nessuna previsione, nessun alert anticipato, riordini d\'emergenza costosi.')

add_solution(doc,
    'Sistema di forecast a due livelli: WMA (Weighted Moving Average) per previsioni base '
    'in tempo reale + Claude AI per analisi approfondita on-demand con rischi e raccomandazioni.')

add_heading(doc, 'Livello 1: Forecast base (WMA)', 2)

add_body(doc,
    'Il sistema analizza i movimenti OUTBOUND degli ultimi 6 mesi, '
    'applica una media mobile ponderata (pesi crescenti per i mesi piu recenti) '
    'e proietta il consumo per i prossimi 3 mesi.')

create_styled_table(doc, ['Dato calcolato', 'Formula', 'Esempio'], [
    ['Consumo mensile (WMA)', 'Media ponderata ultimi 6 mesi', '143 pezzi/mese'],
    ['Stock attuale', 'Somma lotti AVAILABLE', '3.500 pezzi'],
    ['Giorni rimanenti', '(Stock / Consumo mensile) * 30', '733 giorni'],
    ['Riordino necessario', 'Stock <= min_stock OR giorni <= 30', 'Si/No'],
    ['Proiezione 3 mesi', 'Consumo WMA * 3', '143, 143, 143'],
])

add_heading(doc, 'Livello 2: Analisi AI (on-demand)', 2)

add_body(doc,
    'Cliccando "Analisi AI" nel tab Forecast del dettaglio materiale, '
    'Claude riceve i dati del forecast base e restituisce:')

add_bullet(doc, 'Proiezione raffinata per 3 mesi (puo differire dal WMA)', 'Projected:')
add_bullet(doc, 'Punteggio di affidabilita della previsione (0-100%)', 'Confidence:')
add_bullet(doc, 'Spiegazione in italiano di cosa sta succedendo', 'Reasoning:')
add_bullet(doc, 'Fattori di rischio identificati (stagionalita, trend, supply chain)', 'Risks:')

add_heading(doc, 'Alert automatici di riordino', 2)

add_body(doc,
    'Un job periodico (check) analizza tutti i materiali con scorta minima configurata '
    'e crea alert quando lo stock scende sotto la soglia:')

create_styled_table(doc, ['Tipo alert', 'Condizione', 'Azione suggerita'], [
    ['OUT_OF_STOCK', 'Stock = 0', 'Riordino urgente'],
    ['LOW_STOCK', 'Stock <= 50% della scorta minima', 'Riordino prioritario'],
    ['REORDER_SUGGESTED', 'Stock <= scorta minima', 'Riordino consigliato con quantita suggerita'],
])

add_body(doc,
    'Gli alert appaiono come banner nella pagina Materiali e possono essere dismessi '
    'o risolti creando una richiesta di acquisto direttamente dal banner.')

add_use_case(doc, 'Test — Forecast e alert riordino', [
    'Vai in Magazzino > Materiali > clicca su "Cavo RG59"',
    'Vai sul tab "Forecast AI"',
    'Verifica: stock attuale, giorni rimanenti, stato riordino, proiezione 3 mesi',
    'Clicca "Analisi AI" — Claude restituisce previsione con confidence e reasoning',
    'Torna alla lista materiali — se ci sono alert attivi, appaiono come banner arancione/rosso',
    'Clicca "Crea richiesta" dal banner — viene pre-compilata con materiale e quantita suggerita',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 19. OCR FATTURE
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '19. AI — OCR Fatture (Vision Parser)', 1)

add_pain_point(doc,
    'Molte fatture arrivano come PDF non strutturati o immagini scansionate. '
    'Registrarle manualmente nel sistema richiede 15 minuti ciascuna con alto rischio di errori.')

add_solution(doc,
    'Claude Vision analizza PDF e immagini, estrae automaticamente tutti i dati della fattura '
    '(numero, data, importi, righe, P.IVA) e restituisce un punteggio di confidenza per ogni campo.')

add_heading(doc, 'Dati estratti', 2)

create_styled_table(doc, ['Campo', 'Esempio estratto', 'Confidenza tipica'], [
    ['Numero fattura', 'FT-2026-001234', '95-99%'],
    ['Data emissione', '15/03/2026', '98%'],
    ['P.IVA emittente', 'IT12345678901', '97%'],
    ['Importo totale', '2.945,00 EUR', '95%'],
    ['Aliquota IVA', '22%', '93%'],
    ['Righe fattura', 'Monitor HP 24" x 5 @ 589.00', '88-92%'],
    ['Riferimento ordine', 'PR-2026-00042', '90%'],
])

add_heading(doc, 'Flusso', 2)
add_body(doc,
    '1. L\'utente carica un PDF/immagine via upload  |  '
    '2. Claude Vision analizza il documento  |  '
    '3. I campi vengono pre-compilati con punteggio di confidenza  |  '
    '4. L\'utente verifica e conferma  |  '
    '5. La fattura entra nel flusso di matching automatico')

add_use_case(doc, 'Test — Upload fattura con OCR', [
    'Vai in Fatture > Upload Fattura',
    'Carica un file PDF di una fattura',
    'Claude Vision analizza il documento (2-5 secondi)',
    'I campi vengono pre-compilati: numero, data, importo, fornitore',
    'Ogni campo mostra il punteggio di confidenza (verde >90%, arancione 70-90%, rosso <70%)',
    'Correggi eventuali campi con bassa confidenza e conferma',
    'La fattura entra nel sistema e il matching automatico cerca l\'ordine corrispondente',
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 20. CLASSIFICAZIONE EMAIL
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '20. AI — Classificazione Email', 1)

add_pain_point(doc,
    'Le email dai fornitori contengono informazioni critiche (conferme d\'ordine, '
    'aggiornamenti spedizione, fatture) mescolate a email irrilevanti. '
    'Leggerle tutte manualmente richiede tempo e le informazioni importanti si perdono.')

add_solution(doc,
    'Claude classifica automaticamente ogni email in ingresso, '
    'identifica il tipo (conferma ordine, aggiornamento spedizione, offerta, spam) '
    'e suggerisce le azioni da intraprendere.')

add_heading(doc, 'Categorie di classificazione', 2)

create_styled_table(doc, ['Categoria', 'Descrizione', 'Azione automatica'], [
    ['ORDER_CONFIRMATION', 'Conferma d\'ordine dal fornitore', 'Aggiorna stato PR a ORDERED'],
    ['SHIPPING_UPDATE', 'Aggiornamento spedizione/tracking', 'Aggiorna tracking number sulla PR'],
    ['INVOICE', 'Fattura o pro-forma allegata', 'Indirizza al modulo fatturazione'],
    ['QUOTE', 'Preventivo o offerta commerciale', 'Notifica al richiedente'],
    ['GENERAL', 'Comunicazione generica', 'Nessuna azione automatica'],
    ['SPAM', 'Email irrilevante o promozionale', 'Ignora'],
])

add_body(doc,
    'La classificazione avviene nel workflow n8n: l\'email viene ricevuta via IMAP, '
    'il corpo viene inviato a Claude per classificazione, poi il risultato guida '
    'le azioni automatiche nel sistema.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 21. AUTOMAZIONE N8N
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '21. Automazione n8n — Workflow Esterni', 1)

add_body(doc,
    'n8n e una piattaforma di automazione self-hosted che gestisce i workflow esterni di ProcureFlow. '
    'Comunica con l\'applicazione tramite webhook HTTP autenticati con HMAC-SHA256.')

add_pain_point(doc,
    'Senza automazione, ogni email va letta manualmente, ogni fattura registrata a mano, '
    'ogni ritardo scoperto per caso. Il personale spende ore in attivita ripetitive.')

add_solution(doc,
    'n8n automatizza 4 flussi critici: email ingestion, approvazioni, '
    'monitoraggio consegne e sincronizzazione con i portali fornitori.')

add_heading(doc, 'Workflow 1: Email Ingestion Pipeline', 2)

add_body(doc, 'Polling IMAP ogni 5 minuti, parsing AI, creazione/aggiornamento richieste.')

create_styled_table(doc, ['Step', 'Nodo n8n', 'Cosa fa'], [
    ['1', 'IMAP Polling', 'Legge nuove email dalla casella procurement (ogni 5 min)'],
    ['2', 'Filter', 'Filtra solo email da domini vendor noti'],
    ['3', 'AI Classification', 'Claude classifica il tipo di email (ordine/spedizione/fattura/spam)'],
    ['4', 'AI Extraction', 'Estrae dati strutturati: riferimento ordine, stato, data consegna, importo'],
    ['5', 'HTTP Webhook', 'POST /api/webhooks/email-ingestion con dati estratti'],
    ['6', 'ProcureFlow API', 'Upsert PurchaseRequest + crea TimelineEvent'],
    ['7', 'Notification', 'Se lo stato e cambiato, invia notifica al richiedente'],
    ['8', 'Error Handler', 'Log errore + alert all\'admin in caso di fallimento'],
])

add_heading(doc, 'Workflow 2: Approval Automation', 2)

create_styled_table(doc, ['Step', 'Nodo n8n', 'Cosa fa'], [
    ['1', 'Webhook', 'Riceve notifica di nuova richiesta submitted'],
    ['2', 'Rules Engine', 'Valuta soglie: <500 auto, <5000 manager, >=5000 direttore'],
    ['3', 'Auto-Approve', 'Se <500 EUR, approva automaticamente e chiude'],
    ['4', 'Notification', 'Invia email/Slack all\'approvatore con deep link'],
    ['5', 'Wait', 'Attende webhook callback dalla UI (approvazione/rifiuto)'],
    ['6', 'Update', 'POST /api/webhooks/approval-response con decisione'],
])

add_heading(doc, 'Workflow 3: Delivery Monitoring', 2)

create_styled_table(doc, ['Step', 'Nodo n8n', 'Cosa fa'], [
    ['1', 'Cron', 'Esecuzione giornaliera alle 8:00'],
    ['2', 'Query', 'Richieste ORDERED/SHIPPED con consegna prevista <= oggi+3'],
    ['3', 'Overdue Check', 'Se scaduta: cambia priorita a URGENT, notifica richiedente + manager'],
    ['4', 'Reminder', 'Se scade tra 3 giorni: reminder notification'],
    ['5', 'Vendor API', 'Se il vendor ha API: check status automatico'],
    ['6', 'Weekly Report', 'Ogni lunedi: report richieste aperte per department'],
])

add_heading(doc, 'Workflow 4: Vendor Portal Sync', 2)

create_styled_table(doc, ['Step', 'Nodo n8n', 'Cosa fa'], [
    ['1', 'Cron', 'Ogni 30 minuti'],
    ['2', 'Per vendor API', 'Fetch stato ordini dal portale fornitore'],
    ['3', 'Compare', 'Confronta stato attuale vs stato in database'],
    ['4', 'If changed', 'Aggiorna DB + crea TimelineEvent + notifica'],
    ['5', 'Rate Limiter', 'Max 60 richieste/min per vendor'],
])

add_heading(doc, 'Webhook endpoints', 2)

create_styled_table(doc, ['Endpoint', 'Metodo', 'Autenticazione', 'Descrizione'], [
    ['/api/webhooks/email-ingestion', 'POST', 'HMAC-SHA256', 'Riceve dati email parsati'],
    ['/api/webhooks/approval-response', 'POST', 'HMAC-SHA256', 'Callback approvazione'],
    ['/api/webhooks/vendor-update', 'POST', 'HMAC-SHA256', 'Status update da vendor'],
    ['/api/webhooks/sdi-invoice', 'POST', 'HMAC-SHA256', 'Fattura elettronica SDI'],
])

add_body(doc,
    'Ogni webhook e protetto con firma HMAC-SHA256: n8n firma il payload con il segreto '
    'condiviso (N8N_WEBHOOK_SECRET in .env.local), ProcureFlow verifica la firma prima '
    'di processare la richiesta. Richieste con firma non valida vengono rifiutate con 401.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 22. SICUREZZA
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '22. Sicurezza e Autenticazione', 1)

add_heading(doc, 'Autenticazione', 2)

create_styled_table(doc, ['Componente', 'Dettaglio'], [
    ['Provider', 'NextAuth.js con strategia JWT'],
    ['Login', 'Email + password con hash bcrypt'],
    ['MFA', 'TOTP (Time-based One-Time Password) opzionale'],
    ['Session', 'JWT con refresh token automatico'],
    ['Protezione route', 'Middleware su tutte le /api/* e /(dashboard)/*'],
])

add_heading(doc, 'RBAC (Role-Based Access Control)', 2)

create_styled_table(doc, ['Ruolo', 'Richieste', 'Fornitori', 'Approvazioni', 'Utenti', 'Config'], [
    ['ADMIN', 'Tutte', 'CRUD', 'Tutte le soglie', 'CRUD', 'Tutto'],
    ['MANAGER', 'Tutte', 'CRUD', 'Fino a 5.000 EUR', 'Lettura', 'Lettura'],
    ['REQUESTER', 'Solo proprie', 'Lettura', 'Nessuna', 'Nessuno', 'Nessuno'],
    ['VIEWER', 'Lettura', 'Lettura', 'Nessuna', 'Nessuno', 'Nessuno'],
])

add_heading(doc, 'Sicurezza API', 2)
add_bullet(doc, 'Ogni endpoint verifica il JWT e il ruolo dell\'utente', 'Auth check:')
add_bullet(doc, 'Schema Zod su ogni input prima dell\'elaborazione', 'Validazione:')
add_bullet(doc, '100 richieste/minuto per utente', 'Rate limiting:')
add_bullet(doc, 'Webhook protetti con firma HMAC-SHA256', 'HMAC:')
add_bullet(doc, 'Verifica modulo attivo prima di servire le API', 'Module guard:')
add_bullet(doc, 'Max 10 MB, solo PDF/DOCX/XLSX/IMG', 'File upload:')
add_bullet(doc, 'Ogni azione CRUD loggata con autore + timestamp', 'Audit log:')
add_bullet(doc, 'Soft delete, data export, retention policy', 'GDPR:')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 23. ONBOARDING
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '23. Onboarding e Configurazione Iniziale', 1)

add_body(doc,
    'ProcureFlow include un wizard di onboarding per guidare il primo setup dell\'azienda:')

add_bullet(doc, 'Nome azienda, P.IVA, indirizzo', 'Step 1 — Dati azienda:')
add_bullet(doc, 'Quali moduli attivare (Core + opzionali)', 'Step 2 — Moduli:')
add_bullet(doc, 'Import da CSV o creazione manuale', 'Step 3 — Fornitori iniziali:')
add_bullet(doc, 'Invito via email per i colleghi', 'Step 4 — Team:')
add_bullet(doc, 'Riepilogo e conferma configurazione', 'Step 5 — Conferma:')

add_heading(doc, 'Credenziali di test', 2)

create_styled_table(doc, ['Email', 'Password', 'Ruolo'], [
    ['admin@procureflow.it', 'password123', 'ADMIN'],
    ['mario.rossi@saiflow.com', 'password123', 'REQUESTER'],
    ['laura.bianchi@saiflow.com', 'password123', 'MANAGER'],
])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 24. GLOSSARIO
# ════════════════════════════════════════════════════════════════════

add_heading(doc, '24. Glossario Tecnico', 1)

create_styled_table(doc, ['Termine', 'Definizione'], [
    ['PR (Purchase Request)', 'Richiesta di acquisto — l\'unita base del sistema'],
    ['SDI', 'Sistema di Interscambio — infrastruttura pubblica per fatturazione elettronica'],
    ['FatturaPA', 'Formato XML standard per fatture elettroniche italiane'],
    ['Three-Way Matching', 'Riconciliazione a 3 vie: ordine vs fattura vs merce ricevuta'],
    ['WMA', 'Weighted Moving Average — media mobile ponderata per previsioni'],
    ['RBAC', 'Role-Based Access Control — controllo accessi basato su ruoli'],
    ['HMAC', 'Hash-based Message Authentication Code — firma digitale per webhook'],
    ['SSE', 'Server-Sent Events — streaming unidirezionale server>client'],
    ['CIG', 'Codice Identificativo Gara — obbligatorio per acquisti PA'],
    ['CUP', 'Codice Unico Progetto — identificativo progetto pubblico'],
    ['MePA', 'Mercato Elettronico della PA — piattaforma acquisti Consip'],
    ['n8n', 'Piattaforma di automazione workflow self-hosted'],
    ['PgBouncer', 'Connection pooler per PostgreSQL'],
    ['Prisma', 'ORM type-safe per Node.js/TypeScript'],
    ['TanStack Query', 'Libreria di data fetching/caching per React'],
    ['Deploy Config', 'Record database che controlla i moduli attivi'],
    ['Tool Calling', 'Capacita dell\'AI di invocare funzioni per accedere a dati reali'],
    ['Burn Rate', 'Velocita di consumo del budget (EUR/giorno)'],
    ['OCR', 'Optical Character Recognition — riconoscimento ottico caratteri'],
    ['TOTP', 'Time-based One-Time Password — autenticazione a due fattori'],
])

# ── FOOTER ──────────────────────────────────────────────────────────

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ProcureFlow — SaiFlow Hub Centralizzato')
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True

# ── SAVE ────────────────────────────────────────────────────────────

output_path = '/Users/kiraah/Downloads/SaiFlow Hub Centralizzato/procureflow/docs/PROCUREFLOW-GUIDA-COMPLETA.docx'
doc.save(output_path)
print(f'Documento generato: {output_path}')
