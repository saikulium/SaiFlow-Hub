#!/usr/bin/env python3
"""
Generate ProcureFlow Setup Guide — Well-formatted Word document
Parses SETUP-GUIDE.md and produces a professional .docx
"""

import re
import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ──────────────────────────────────────────────────────────
INDIGO = RGBColor(0x63, 0x66, 0xF1)
INDIGO_LIGHT = RGBColor(0x81, 0x8C, 0xF8)
DARK = RGBColor(0x1C, 0x1C, 0x1F)
GRAY = RGBColor(0x71, 0x71, 0x7A)
LIGHT_GRAY = RGBColor(0xA1, 0xA1, 0xAA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xEF, 0x44, 0x44)
GREEN = RGBColor(0x22, 0xC5, 0x5E)
ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
BLUE = RGBColor(0x3B, 0x82, 0xF6)
CODE_BG = 'F4F4F5'
CALLOUT_BG = 'EEF2FF'
TABLE_HEADER_BG = '6366F1'
TABLE_ALT_BG = 'F9FAFB'


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)


def set_paragraph_border_left(paragraph, color_hex='6366F1', width='12'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), width)
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=60):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=10, font='Calibri'):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return run


def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 5):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            hs = doc.styles[style_name]
            hs.font.name = 'Calibri'
            hs.font.color.rgb = INDIGO if level <= 2 else DARK
            if level == 1:
                hs.font.size = Pt(22)
                hs.paragraph_format.space_before = Pt(24)
                hs.paragraph_format.space_after = Pt(8)
            elif level == 2:
                hs.font.size = Pt(16)
                hs.paragraph_format.space_before = Pt(18)
                hs.paragraph_format.space_after = Pt(6)
            elif level == 3:
                hs.font.size = Pt(13)
                hs.paragraph_format.space_before = Pt(14)
                hs.paragraph_format.space_after = Pt(4)
            else:
                hs.font.size = Pt(11)


def add_code_block(doc, code_text):
    """Add a styled code block with gray background."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        set_paragraph_shading(p, CODE_BG)
        set_paragraph_spacing(p, before=0, after=0)
        pf = p.paragraph_format
        pf.left_indent = Cm(0.5)
        pf.right_indent = Cm(0.5)
        add_run(p, line, size=8.5, font='Consolas', color=DARK)


def add_callout(doc, text, color_hex=CALLOUT_BG, border_color='6366F1'):
    """Add a callout/blockquote box."""
    p = doc.add_paragraph()
    set_paragraph_shading(p, color_hex)
    set_paragraph_border_left(p, border_color)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    # Parse bold markers in text
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if i % 2 == 1:
            add_run(p, part, bold=True, size=9, color=DARK)
        else:
            add_run(p, part, size=9, color=GRAY)


def add_styled_table(doc, headers, rows):
    """Create a well-formatted table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = WHITE
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            p = cell.paragraphs[0]
            # Parse bold/code from markdown
            parts = re.split(r'\*\*(.+?)\*\*|`(.+?)`', cell_text)
            for j, part in enumerate(parts):
                if part is None:
                    continue
                if j % 3 == 1:  # bold
                    run = p.add_run(part)
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
                elif j % 3 == 2:  # code
                    run = p.add_run(part)
                    run.font.size = Pt(8.5)
                    run.font.name = 'Consolas'
                    run.font.color.rgb = INDIGO
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
            if row_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    doc.add_paragraph()  # spacing after table


# ── Markdown Parser ─────────────────────────────────────────────────

def parse_md_line_rich(doc, text, base_size=10):
    """Add a paragraph with inline markdown formatting (bold, code, italic)."""
    p = doc.add_paragraph()
    # Split on **bold**, `code`, *italic*
    pattern = r'\*\*(.+?)\*\*|`(.+?)`|\*(.+?)\*'
    parts = re.split(pattern, text)
    for i, part in enumerate(parts):
        if part is None:
            continue
        mod = i % 4
        if mod == 1:  # bold
            add_run(p, part, bold=True, size=base_size, color=DARK)
        elif mod == 2:  # code
            add_run(p, part, size=base_size - 1, font='Consolas', color=INDIGO)
        elif mod == 3:  # italic
            add_run(p, part, italic=True, size=base_size, color=GRAY)
        else:
            add_run(p, part, size=base_size, color=DARK)
    return p


def parse_table(lines):
    """Parse markdown table lines into headers and rows."""
    headers = []
    rows = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if i == 0:
            headers = cells
        elif set(c.strip() for c in cells) <= {'-', '---', '----', '-----', '------', '-------', '--------', '---------', '----------'}:
            continue  # separator
        elif all(re.match(r'^-+$', c.strip()) for c in cells):
            continue
        else:
            rows.append(cells)
    return headers, rows


def convert_md_to_docx(md_path, docx_path):
    doc = Document()
    setup_styles(doc)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip markdown HRs
        if stripped == '---':
            # Add a subtle line break
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=6)
            i += 1
            continue

        # Code block start/end
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # H1
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:].strip()
            # Cover page
            doc.add_paragraph()  # spacing
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, 'PROCUREFLOW', bold=True, size=12, color=INDIGO)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, title_text, bold=True, size=26, color=DARK)
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, f'Versione {datetime.date.today().strftime("%d/%m/%Y")}', size=10, color=GRAY)
            doc.add_page_break()
            i += 1
            continue

        # H2
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # H3
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # H4
        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            doc.add_heading(text, level=4)
            i += 1
            continue

        # Table detection
        if stripped.startswith('|') and i + 1 < len(lines):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            headers, rows = parse_table(table_lines)
            if headers and rows:
                add_styled_table(doc, headers, rows)
            continue

        # Blockquote / callout
        if stripped.startswith('> '):
            callout_text = stripped[2:].strip()
            # Collect multi-line blockquotes
            while i + 1 < len(lines) and lines[i + 1].strip().startswith('> '):
                i += 1
                next_text = lines[i].strip()[2:].strip()
                callout_text += ' ' + next_text
            # Remove markdown link syntax
            callout_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', callout_text)
            add_callout(doc, callout_text)
            i += 1
            continue

        # Bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # strip links
            p = parse_md_line_rich(doc, text, base_size=9.5)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            pf.first_line_indent = Cm(-0.4)
            # Add bullet character
            if p.runs:
                p.runs[0].text = '\u2022  ' + p.runs[0].text
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2).strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = parse_md_line_rich(doc, text, base_size=9.5)
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            pf.first_line_indent = Cm(-0.5)
            if p.runs:
                p.runs[0].text = f'{num}.  ' + p.runs[0].text
            i += 1
            continue

        # Indented sub-items (  - text or   1. text)
        indent_match = re.match(r'^(\s{2,})([-*]|\d+\.)\s+(.+)', line)
        if indent_match:
            prefix = indent_match.group(2)
            text = indent_match.group(3).strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = parse_md_line_rich(doc, text, base_size=9)
            pf = p.paragraph_format
            pf.left_indent = Cm(1.8)
            pf.first_line_indent = Cm(-0.4)
            bullet = '\u2013  ' if prefix in ['-', '*'] else f'{prefix} '
            if p.runs:
                p.runs[0].text = bullet + p.runs[0].text
            i += 1
            continue

        # TOC lines (skip them - we generate our own structure)
        if re.match(r'^\d+\.\s+\[', stripped):
            i += 1
            continue
        if re.match(r'^\s+-\s+\[', stripped):
            i += 1
            continue

        # Regular paragraph
        text = stripped
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # strip links
        parse_md_line_rich(doc, text)
        i += 1

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '— ProcureFlow —', size=9, color=LIGHT_GRAY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f'Documento generato il {datetime.date.today().strftime("%d/%m/%Y")}', size=8, color=LIGHT_GRAY)

    doc.save(docx_path)
    print(f'Documento generato: {docx_path}')


if __name__ == '__main__':
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(script_dir, 'SETUP-GUIDE.md')
    docx_path = os.path.join(script_dir, 'ProcureFlow-Setup-Guide.docx')
    convert_md_to_docx(md_path, docx_path)
